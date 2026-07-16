from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/AgentScope-2.0-用量与运行记录实施说明.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=11, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        set_cell_shading(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=9, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=9, color="334155")
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("AgentScope DataAgent | Internal Implementation Note")
    set_run_font(run, size=8, color="64748B")
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_p.add_run("2026-07-13")
    set_run_font(run, size=8, color="64748B")


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("AgentScope 2.0 用量、运行记录与租户模型实施说明")
    set_run_font(run, size=22, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("适用范围：agentscope-dataagent | 版本：2.0.0-SNAPSHOT | 更新：2026-07-13")
    set_run_font(run, size=10, color="64748B")

    callout = doc.add_table(rows=1, cols=1)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("结论：请求级用量、独立运行记录和租户模型连接均由 AgentScope 2.0 的稳定扩展点实现；应用层不接管 Agent、Sandbox 或 span 生命周期。")
    set_run_font(run, size=10, color=INK, bold=True)
    set_table_geometry(callout, [9360])

    doc.add_heading("1. 目标与范围", level=1)
    add_paragraph(doc, "本次实现将原先仅在内存中保存的“聊天轮次计数”替换为 MySQL 请求级账本，并接入 AgentScope 的 ChatUsage、OtelTracingMiddleware 和 ModelCreationContext。目标是让产品同时具备可计量、可排障、可按租户配置模型的运行基础。")
    add_table(doc, ["能力", "实施结果", "边界"], [
        ["用量", "按一次聊天请求汇总 token、缓存 token、耗时、模型、成本与状态", "成本取决于已配置的模型单价；未配置时安全记为 0"],
        ["运行记录", "独立页面按 trace 展开 Agent、模型和工具 span", "聊天气泡内的执行轨迹继续保留，不迁移、不替代"],
        ["租户模型", "个人 Agent 用 ModelCreationContext 读取加密凭据和 Base URL", "当前用户 ID 是过渡 tenant ID；全局 Agent 仍用静态模型"],
    ], [1875, 3385, 4100])

    doc.add_heading("2. 架构决策", level=1)
    add_paragraph(doc, "框架负责运行时生命周期，应用负责业务索引和治理。这个分层避免再次引入应用自管沙箱或自定义 tracing 链路。")
    add_bullet(doc, "ChatUsage：从每个 ModelCallEndEvent 收集 input、output 与 cached prompt tokens，并按请求汇总。")
    add_bullet(doc, "OtelTracingMiddleware：附加到每个 HarnessAgent Builder，生成 Agent、模型、工具三个层次的 span。")
    add_bullet(doc, "OpenTelemetry SDK：使用 MySQL span exporter 持久化白名单属性；提示词和模型正文不进入运行记录索引。")
    add_bullet(doc, "ModelCreationContext：用 tenant-openai / tenant-ollama factory 依据加密凭据、Base URL、endpoint 和缓存 ID 构建模型。")
    add_bullet(doc, "LongCat：保持 OpenAI-compatible formatter fallback；不强开 native JSON Schema 或 native structured output。")

    doc.add_heading("3. 数据模型与成本口径", level=1)
    add_table(doc, ["表", "关键字段", "用途"], [
        ["usage_event", "tenant_id, user_id, agent_id, session_key, model_id, input_tokens, output_tokens, cached_prompt_tokens, duration_ms, cost_microusd, outcome", "一次聊天请求的聚合账本"],
        ["trace_run", "trace_id, user_id, agent_id, session_key, model_id, status, started_at_ms, ended_at_ms", "运行记录页的查询根节点"],
        ["trace_span", "trace_id, span_id, parent_span_id, operation_name, status, duration_ms, attributes_json", "AgentScope OTel span 的安全索引"],
        ["tenant_model_config", "tenant_id, logical_model_id, provider, model_name, base_url, api_key_ciphertext, price fields", "租户模型连接与价格策略"],
    ], [1875, 4750, 2735])
    add_paragraph(doc, "成本统一使用 micro USD 整数存储，避免浮点误差。标准输入 token、缓存输入 token 与输出 token 分别按“每百万 token 的 micro USD 单价”计算；标准输入为 max(input - cached, 0)。")

    doc.add_heading("4. 请求执行与状态", level=1)
    add_number(doc, "聊天控制器创建根 span 和 trace_run，并将 span context 放入 Reactor 流。")
    add_number(doc, "AgentScope OtelTracingMiddleware 在同一 context 下创建 Agent、模型和工具子 span。")
    add_number(doc, "流式事件中的 ModelCallEndEvent 累加 ChatUsage；工具事件继续推送至聊天内的折叠执行轨迹。")
    add_number(doc, "请求结束时写入 usage_event，计算成本，结束根 span 并更新 trace_run 为 SUCCESS、CANCELLED 或 ERROR。")
    add_number(doc, "span exporter 只导出模型名、token、工具名、调用数、reply ID 等白名单字段，运行记录页按 trace 查询展示。")

    doc.add_heading("5. 前端交互", level=1)
    add_table(doc, ["页面", "用途", "关键交互"], [
        ["/usage", "个人用量", "展示总 token、缓存 token、成本、平均耗时、请求趋势和模型分布"],
        ["/traces", "运行记录", "一条请求一行，可展开 span、耗时、状态和安全属性"],
        ["/models", "模型连接", "写入 provider、模型名、Base URL、密钥和单价；密钥只显示已配置状态"],
        ["/chat", "对话", "继续保留消息内工具调用、人工确认和执行轨迹"],
    ], [1500, 2300, 5560])

    doc.add_heading("6. 租户模型边界与安全要求", level=1)
    add_paragraph(doc, "当前没有独立组织租户表，因此 tenantForUser(userId) 是明确的过渡策略：数据模型保留 tenant_id，未来引入组织后只替换这一映射，不迁移模型配置表。")
    add_bullet(doc, "API Key 通过 AES-GCM 加密后写入 MySQL，接口只返回 apiKeyConfigured，不回显原文。")
    add_bullet(doc, "生产环境必须设置 DATAAGENT_CREDENTIAL_ENCRYPTION_KEY；更换该值会使已有密钥无法解密。")
    add_bullet(doc, "租户连接变更会失效当前用户的个人 Agent 缓存，新的个人 Agent 重建后使用新配置。")
    add_bullet(doc, "没有租户配置时继续解析现有 static local / longcat；全局 Data Agent 也保持现有静态注册，以控制迁移风险。")

    doc.add_heading("7. 验收与运行手册", level=1)
    add_number(doc, "完成一段含模型和工具调用的对话，确认 /usage 出现 token、缓存 token、模型、耗时和成本记录。")
    add_number(doc, "打开 /traces，展开最新运行记录，确认 Agent、模型、工具 span 串在同一 trace 下。")
    add_number(doc, "点击聊天停止，确认本次 trace 和 usage 的 outcome 为 CANCELLED。")
    add_number(doc, "在 /models 新增 OpenAI-compatible 连接并保存，确认密钥不回显；新建或重建个人 Agent 后验证新连接。")
    add_number(doc, "执行 mvn test 与 frontend 下 npm run build。当前验收基线为后端 40 通过、11 个既有跳过，以及 TypeScript/Vite 生产构建通过。")

    doc.add_heading("8. 后续建议", level=1)
    add_bullet(doc, "引入正式 tenant/organization 领域模型后，将 tenantForUser 映射替换为成员归属解析，并补充管理员级跨租户用量和 trace 查询。")
    add_bullet(doc, "为 trace_run 与 usage_event 增加保留策略和分区/归档策略，避免长期高频请求导致主库膨胀。")
    add_bullet(doc, "在确认某个兼容端点完整支持 tools + JSON Schema 后，再以显式 capability flag 逐模型启用 native structured output。")

    doc.save(OUT)


if __name__ == "__main__":
    main()
